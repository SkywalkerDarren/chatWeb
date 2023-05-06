import os
import time

import PyPDF2
import docx
import readability
from langdetect import detect
from newspaper import fulltext, Article
from selenium import webdriver


def web_crawler_newspaper(url: str) -> tuple[list[str], str]:
    """Run the web crawler."""
    raw_html, lang = _get_raw_html(url)
    try:
        text = fulltext(raw_html, language=lang)
    except:
        article = Article(url)
        article.download()
        article.parse()
        text = article.text
    contents = [text.strip() for text in text.splitlines() if text.strip()]
    return contents, lang


def _get_raw_html(url):
    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_argument('--headless')
    chrome_options.add_argument('--disable-gpu')
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                                'AppleWebKit/537.36 (KHTML, like Gecko) Chrome/112.0.0.0 Safari/537.36')

    with webdriver.Chrome(options=chrome_options) as driver:
        driver.get(url)
        print("Please wait for 5 seconds until the webpage finishes loading.")
        time.sleep(5)
        html = driver.page_source

    doc = readability.Document(html)
    html = doc.summary()
    lang = detect(html)
    return html, lang[0:2]


def extract_text_from_pdf(file_path: str) -> tuple[list[str], str]:
    """Extract text content from a PDF file."""
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        contents = []
        for page in pdf_reader.pages:
            page_text = page.extract_text().strip()
            raw_text = [text.strip() for text in page_text.splitlines() if text.strip()]
            new_text = ''
            for text in raw_text:
                new_text += text
                if text[-1] in ['.', '!', '?', '。', '！', '？', '…', ';', '；', ':', '：', '”', '’', '）', '】', '》', '」',
                                '』', '〕', '〉', '》', '〗', '〞', '〟', '»', '"', "'", ')', ']', '}']:
                    contents.append(new_text)
                    new_text = ''
            if new_text:
                contents.append(new_text)
        lang = detect('\n'.join(contents))
        return contents, lang[0:2]


def extract_text_from_txt(file_path: str) -> tuple[list[str], str]:
    """Extract text content from a TXT file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        contents = [text.strip() for text in f.readlines() if text.strip()]
        lang = detect('\n'.join(contents))
        return contents, lang[0:2]


def extract_text_from_docx(file_path: str) -> tuple[list[str], str]:
    """Extract text content from a DOCX file."""
    document = docx.Document(file_path)
    contents = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    lang = detect('\n'.join(contents))
    return contents, lang[0:2]
